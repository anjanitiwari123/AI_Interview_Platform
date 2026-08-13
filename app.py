from __future__ import annotations

import tempfile
import uuid
from pathlib import Path

import streamlit as st

from audio_analysis import (
    analyze_audio,
    save_recorded_audio,
    transcribe_audio,
)

from feedback import generate_interview_feedback

from nlp_analysis import (
    evaluate_live_answer,
    generate_interview_question,
)

from report import (
    build_interview_report,
    save_json_report,
    save_pdf_report,
)


BASE_DIRECTORY = Path(__file__).parent
REPORT_DIRECTORY = BASE_DIRECTORY / "reports"


def clear_interview_state():
    keys = [
        "interview_started",
        "interview_complete",
        "current_question",
        "question_results",
        "interview_id",
        "report_data",
        "report_paths",
        "resume_text",
        "resume_name",
        "pending_transcript",
        "pending_audio_metrics",
    ]

    for key in keys:
        st.session_state.pop(key, None)



def extract_resume_text(uploaded_resume):
    """Read text from PDF, DOCX, or TXT resume."""

    filename = uploaded_resume.name.lower()

    if filename.endswith(".pdf"):

        try:
            from pypdf import PdfReader

        except ImportError as error:
            raise RuntimeError(
                "Install pypdf: pip install pypdf"
            ) from error

        reader = PdfReader(uploaded_resume)

        text = "\n".join(
            page.extract_text() or ""
            for page in reader.pages
        )


    elif filename.endswith(".docx"):

        try:
            from docx import Document

        except ImportError as error:
            raise RuntimeError(
                "Install python-docx"
            ) from error

        document = Document(uploaded_resume)

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
        ]

        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]

        text = "\n".join(
            paragraphs + table_cells
        )


    elif filename.endswith(".txt"):

        text = uploaded_resume.getvalue().decode(
            "utf-8",
            errors="ignore"
        )

    else:
        raise ValueError(
            "Upload PDF, DOCX, or TXT resume."
        )


    cleaned_text = "\n".join(
        line.strip()
        for line in text.splitlines()
        if line.strip()
    )


    if len(cleaned_text) < 80:
        raise ValueError(
            "Resume text is too short."
        )


    return cleaned_text



def start_interview(
        topic,
        question_count,
        resume_text="",
        resume_name=""
):

    first_question = generate_interview_question(
        topic,
        [],
        resume_text
    )


    st.session_state.interview_started = True
    st.session_state.interview_complete = False

    st.session_state.current_question = first_question

    st.session_state.question_results = []

    st.session_state.interview_id = uuid.uuid4().hex

    st.session_state.interview_topic = (
        topic
        or
        "Resume-based technical interview"
    )

    st.session_state.question_count = question_count

    st.session_state.resume_text = resume_text

    st.session_state.resume_name = resume_name



def show_completed_question(result, number):

    analysis = result["technical_analysis"]

    with st.expander(
        f"Question {number}: {analysis['technical_score']}/10",
        expanded=False
    ):

        st.write(
            analysis["question"]
        )

        st.write(
            "**Transcript:**",
            result["transcript"]
        )

        st.write(
            "**Covered points:**",
            ", ".join(
                analysis["covered_points"]
            )
            or
            "None"
        )

        st.write(
            "**Missing points:**",
            ", ".join(
                analysis["missing_points"]
            )
            or
            "None"
        )

        st.write(
            "**Feedback:**",
            analysis["feedback"]
        )



def submit_answer(
        recorded_audio,
        whisper_model
):

    """
    Convert recorded audio to text.
    """

    with tempfile.TemporaryDirectory() as temp_dir:

        audio_path = Path(temp_dir) / "live_answer.wav"


        try:

            with st.spinner(
                "Converting live speech to text..."
            ):

                save_recorded_audio(
                    recorded_audio.getvalue(),
                    audio_path
                )


                transcript = transcribe_audio(
                    audio_path,
                    whisper_model
                )


                audio_metrics = analyze_audio(
                    audio_path,
                    transcript
                )


        except Exception as error:

            st.exception(error)
            return



    st.session_state.pending_transcript = transcript

    st.session_state.pending_audio_metrics = audio_metrics



def _complete_answer(
        transcript,
        audio_metrics
):

    with st.spinner(
        "Evaluating the answer..."
    ):


        technical_analysis = evaluate_live_answer(

            st.session_state.current_question,

            transcript

        )


    st.session_state.question_results.append({

        "question_number":
        len(st.session_state.question_results) + 1,


        "transcript":
        transcript,


        "audio_analysis":
        audio_metrics,


        "technical_analysis":
        technical_analysis

    })



    if len(st.session_state.question_results) >= st.session_state.question_count:


        feedback = generate_interview_feedback(
            st.session_state.question_results
        )


        st.session_state.report_data = build_interview_report(

            st.session_state.interview_topic,

            st.session_state.question_results,

            feedback,

            st.session_state.resume_name

        )


        st.session_state.interview_complete = True

        st.session_state.current_question = None

        return



    previous_questions = [

        item["technical_analysis"]["question"]

        for item in st.session_state.question_results

    ]


    try:

        with st.spinner(
            "Generating next question..."
        ):

            st.session_state.current_question = generate_interview_question(

                st.session_state.interview_topic,

                previous_questions,

                st.session_state.resume_text

            )


    except Exception as error:

        st.exception(error)
def evaluate_pending_answer():

    """
    Evaluate edited transcript instead of raw Whisper output.
    """

    transcript = (
        st.session_state.pending_transcript
        .strip()
    )


    if len(transcript.split()) < 3:

        st.warning(
            "Please enter at least a short answer, or record again."
        )

        return



    audio_metrics = dict(
        st.session_state.pending_audio_metrics
    )


    word_count = len(
        [
            word
            for word in transcript.split()
            if any(char.isalnum() for char in word)
        ]
    )


    duration = max(
        float(
            audio_metrics.get(
                "duration_seconds",
                0
            )
        ),
        0.1
    )


    audio_metrics.update({

        "transcript":
        transcript,

        "word_count":
        word_count,

        "words_per_minute":
        round(
            word_count / duration * 60,
            1
        )

    })


    try:

        _complete_answer(
            transcript,
            audio_metrics
        )


        st.session_state.pop(
            "pending_transcript",
            None
        )


        st.session_state.pop(
            "pending_audio_metrics",
            None
        )


    except Exception as error:

        st.exception(error)



def show_pending_review():

    st.subheader(
        "Check transcript before scoring"
    )


    st.caption(
        "Edit Whisper mistakes before evaluation."
    )


    st.text_area(
        "Transcript",
        key="pending_transcript",
        height=180
    )


    score_column, retry_column = st.columns(2)


    if score_column.button(
        "Evaluate reviewed answer",
        type="primary"
    ):

        evaluate_pending_answer()

        if "pending_transcript" not in st.session_state:

            st.rerun()



    if retry_column.button(
        "Record again"
    ):

        st.session_state.pop(
            "pending_transcript",
            None
        )

        st.session_state.pop(
            "pending_audio_metrics",
            None
        )

        st.rerun()



def show_final_report():

    report_data = st.session_state.report_data


    st.success(
        "Interview completed."
    )


    st.metric(
        "Overall technical score",
        f"{report_data['overall_technical_score']}/10"
    )


    st.write(
        "### Strengths"
    )

    for item in report_data["strengths"]:

        st.write(
            "-",
            item
        )


    st.write(
        "### Weak areas"
    )

    for item in report_data["weak_areas"]:

        st.write(
            "-",
            item
        )


    st.write(
        "### Final feedback"
    )


    st.write(
        report_data["final_feedback"]
    )



    if "report_paths" not in st.session_state:

        st.session_state.report_paths = (

            save_json_report(
                report_data,
                REPORT_DIRECTORY
            ),

            save_pdf_report(
                report_data,
                REPORT_DIRECTORY
            )

        )


    json_path, pdf_path = st.session_state.report_paths


    json_column, pdf_column = st.columns(2)


    json_column.download_button(

        "Download JSON Report",

        json_path.read_bytes(),

        json_path.name,

        "application/json"

    )


    pdf_column.download_button(

        "Download PDF Report",

        pdf_path.read_bytes(),

        pdf_path.name,

        "application/pdf"

    )



    if st.button(
        "Start New Interview"
    ):

        clear_interview_state()

        st.rerun()



def main():

    st.set_page_config(

        page_title="Live AI Interview",

        page_icon="🎙️",

        layout="wide"

    )


    st.title(
        "Real-Time AI Interview Practice"
    )


    st.caption(
        "AI interview practice using Speech Recognition, NLP and LLM evaluation."
    )



    with st.sidebar:


        topic = st.text_input(

            "Target role or topic",

            value=""

        )


        uploaded_resume = st.file_uploader(

            "Upload resume",

            type=[
                "pdf",
                "docx",
                "txt"
            ]

        )


        question_count = st.selectbox(

            "Number of questions",

            [1,2,3,4,5],

            index=2

        )


        whisper_model = st.selectbox(

            "Whisper model",

            [
                "tiny",
                "base",
                "small"
            ],

            index=1

        )


        st.caption(
            "LLM evaluation powered by Groq."
        )



    if not st.session_state.get(
        "interview_started"
    ):


        st.write(

            "Upload resume to generate personalized technical questions."

        )


        if st.button(

            "Start Interview",

            type="primary"

        ):


            try:


                resume_text = ""

                resume_name = ""



                if uploaded_resume is not None:


                    with st.spinner(

                        "Reading resume..."

                    ):


                        resume_text = extract_resume_text(

                            uploaded_resume

                        )


                    resume_name = uploaded_resume.name



                with st.spinner(

                    "Generating first question..."

                ):


                    start_interview(

                        topic,

                        question_count,

                        resume_text,

                        resume_name

                    )


                st.rerun()



            except Exception as error:

                st.exception(error)


        return



    if st.session_state.get(
        "interview_complete"
    ):

        show_final_report()

        return



    for index, result in enumerate(

        st.session_state.question_results,

        start=1

    ):

        show_completed_question(
            result,
            index
        )



    current_number = (

        len(
            st.session_state.question_results
        )

        + 1

    )


    question_data = st.session_state.current_question



    st.subheader(

        f"Question {current_number} of {st.session_state.question_count}"

    )


    st.info(
        question_data["question"]
    )


    st.write(

        "**Difficulty:**",

        question_data["difficulty"].title()

    )


    st.write(

        "**Evaluation points:**",

        ", ".join(
            question_data["evaluation_points"]
        )

    )



    try:

        from streamlit_webrtc import (

            WebRtcMode,

            webrtc_streamer

        )


    except ImportError:

        st.error(
            "Install streamlit-webrtc and av."
        )

        return


    RTC_CONFIGURATION = {
    "iceServers": [
        {
            "urls": [
                "stun:stun.l.google.com:19302"
            ]
        }
    ]
    }

    camera_context = webrtc_streamer(

        key=f"interview-camera-{st.session_state.interview_id}",
        mode=WebRtcMode.SENDONLY,
        rtc_configuration=RTC_CONFIGURATION,
        media_stream_constraints={

            "video": True,

            "audio": False

        },

        async_processing=True,

        media_toggle_controls=True

        )


    if not camera_context.state.playing:
        st.caption(
            
            "Start camera permission."
        )



    if "pending_transcript" in st.session_state:

        show_pending_review()

        return



    st.subheader(
        "Record your answer"
    )


    recorded_audio = st.audio_input(

        "Answer recording",

        sample_rate=16000,

        key=f"answer-audio-{st.session_state.interview_id}-{current_number}"

    )


    if recorded_audio is not None:


        st.audio(
            recorded_audio
        )


        if st.button(

            "Transcribe and review answer",

            type="primary"

        ):


            submit_answer(

                recorded_audio,

                whisper_model

            )


            if "pending_transcript" in st.session_state:

                show_pending_review()



if __name__ == "__main__":

    main()
